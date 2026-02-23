"""DOCX import/export using python-docx — with layout metadata extraction."""

from __future__ import annotations

from typing import TYPE_CHECKING, Any

from docx import Document
from docx.table import Table as DocxTable

from uaf.app.formats import ComparisonResult
from uaf.core.edges import Edge, EdgeType
from uaf.core.node_id import EdgeId, NodeId, utc_now
from uaf.core.nodes import (
    Artifact,
    Cell,
    Heading,
    LayoutHint,
    NodeType,
    Paragraph,
    Sheet,
    make_node_metadata,
)

if TYPE_CHECKING:
    from pathlib import Path

    from uaf.db.graph_db import GraphDB

# Conversion factor: EMU (English Metric Units) to points.
_EMU_TO_PT = 1.0 / 12700.0


class DocxHandler:
    """Import/export DOCX files via the UAF graph."""

    def import_file(self, path: Path, db: GraphDB) -> NodeId:
        """Parse a DOCX file into UAF nodes (paragraphs, headings, tables)."""
        doc = Document(str(path))

        # Extract page geometry from the first section.
        art_layout = _extract_page_layout(doc)
        art = Artifact(
            meta=make_node_metadata(NodeType.ARTIFACT, layout=art_layout),
            title=path.stem,
        )
        art_id = db.create_node(art)

        content_width = _content_width_pt(doc)

        for reading_order, item in enumerate(doc.iter_inner_content()):
            if isinstance(item, DocxTable):
                self._import_table(item, art_id, db)
            else:
                self._import_paragraph(
                    item, art_id, db,
                    reading_order=reading_order,
                    content_width=content_width,
                )

        return art_id

    def export_file(self, db: GraphDB, root_id: NodeId, path: Path) -> None:
        """Export a UAF artifact as a DOCX file."""
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()
        children = db.get_children(root_id)

        for child in children:
            if isinstance(child, Heading):
                level = max(0, min(child.level, 9))
                doc.add_heading(child.text, level=level)
            elif isinstance(child, Paragraph):
                p = doc.add_paragraph(child.text)
                if child.style == "body":
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            elif isinstance(child, Sheet):
                self._export_table(db, child, doc)

        doc.save(str(path))

    def _import_paragraph(
        self,
        para: Any,
        parent_id: NodeId,
        db: GraphDB,
        *,
        reading_order: int = 0,
        content_width: float | None = None,
    ) -> None:
        """Convert a python-docx paragraph into a UAF node with layout metadata."""
        text = para.text.strip()
        if not text:
            return

        style_name = (para.style.name if para.style else "").lower()
        layout = _extract_paragraph_layout(
            para,
            reading_order=reading_order,
            content_width=content_width,
        )

        node: Heading | Paragraph
        if style_name.startswith("heading"):
            level = _parse_heading_level(style_name)
            node = Heading(
                meta=make_node_metadata(NodeType.HEADING, layout=layout),
                text=text,
                level=level,
            )
        else:
            node = Paragraph(
                meta=make_node_metadata(NodeType.PARAGRAPH, layout=layout),
                text=text,
            )

        nid = db.create_node(node)
        db.create_edge(_contains(parent_id, nid))

    def _import_table(
        self, table: DocxTable, parent_id: NodeId, db: GraphDB,
    ) -> None:
        """Convert a python-docx table into a UAF Sheet + Cell nodes."""
        rows = table.rows
        num_rows = len(rows)
        num_cols = max((len(row.cells) for row in rows), default=0)

        sheet = Sheet(
            meta=make_node_metadata(NodeType.SHEET),
            title="Table",
            rows=num_rows,
            cols=num_cols,
        )
        sheet_id = db.create_node(sheet)
        db.create_edge(_contains(parent_id, sheet_id))

        for r, row in enumerate(rows):
            for c, cell in enumerate(row.cells):
                cell_node = Cell(
                    meta=make_node_metadata(NodeType.CELL),
                    value=cell.text.strip(),
                    row=r,
                    col=c,
                )
                cell_id = db.create_node(cell_node)
                db.create_edge(_contains(sheet_id, cell_id))

    def _export_table(self, db: GraphDB, sheet: Sheet, doc: Any) -> None:
        """Export a Sheet node as a DOCX table."""
        cells = db.get_children(sheet.meta.id)
        grid: list[list[str]] = [[""] * sheet.cols for _ in range(sheet.rows)]

        for cell in cells:
            if isinstance(cell, Cell) and cell.row < sheet.rows and cell.col < sheet.cols:
                grid[cell.row][cell.col] = (
                    str(cell.value) if cell.value is not None else ""
                )

        table = doc.add_table(rows=sheet.rows, cols=sheet.cols)
        for r, row_data in enumerate(grid):
            for c, value in enumerate(row_data):
                table.rows[r].cells[c].text = value


class DocxComparator:
    """Compare two DOCX files for content equivalence."""

    def compare(self, original: Path, rebuilt: Path) -> ComparisonResult:
        """Compare DOCX files by extracted text content."""
        orig_texts = _extract_texts(original)
        rebuilt_texts = _extract_texts(rebuilt)

        differences: list[str] = []
        ignored: list[str] = ["DOCX formatting/styles may differ"]

        max_len = max(len(orig_texts), len(rebuilt_texts))
        matching = 0

        for i in range(max_len):
            orig = orig_texts[i] if i < len(orig_texts) else ""
            rebuilt_val = rebuilt_texts[i] if i < len(rebuilt_texts) else ""
            if orig == rebuilt_val:
                matching += 1
            else:
                differences.append(f"Block {i}: {orig!r} != {rebuilt_val!r}")

        score = matching / max_len if max_len > 0 else 1.0

        return ComparisonResult(
            is_equivalent=len(differences) == 0,
            differences=differences,
            ignored=ignored,
            similarity_score=score,
        )


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _contains(source: NodeId, target: NodeId) -> Edge:
    return Edge(
        id=EdgeId.generate(),
        source=source,
        target=target,
        edge_type=EdgeType.CONTAINS,
        created_at=utc_now(),
    )


def _parse_heading_level(style_name: str) -> int:
    """Extract heading level from a style name like 'heading 1'."""
    for part in style_name.split():
        if part.isdigit():
            return max(1, min(int(part), 6))
    return 1


def _extract_texts(path: Path) -> list[str]:
    """Extract text content from a DOCX file for comparison."""
    doc = Document(str(path))
    texts: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            texts.append(text)
    return texts


def _extract_page_layout(doc: Any) -> LayoutHint | None:
    """Extract page dimensions from the first DOCX section."""
    try:
        section = doc.sections[0]
    except (IndexError, AttributeError):
        return None

    pw = section.page_width
    ph = section.page_height
    if pw is None or ph is None:
        return None

    return LayoutHint(
        width=float(pw) * _EMU_TO_PT,
        height=float(ph) * _EMU_TO_PT,
    )


def _content_width_pt(doc: Any) -> float | None:
    """Calculate content width (page width minus margins) in points."""
    try:
        section = doc.sections[0]
    except (IndexError, AttributeError):
        return None

    pw = section.page_width
    lm = section.left_margin
    rm = section.right_margin
    if pw is None or lm is None or rm is None:
        return None

    return float(pw - lm - rm) * _EMU_TO_PT


def _extract_paragraph_layout(
    para: Any,
    *,
    reading_order: int = 0,
    content_width: float | None = None,
) -> LayoutHint | None:
    """Extract font and layout metadata from a python-docx paragraph."""
    font_family: str | None = None
    font_size: float | None = None
    font_weight: str | None = None
    font_style: str | None = None

    # Extract font info from the first run that has data.
    for run in getattr(para, "runs", []):
        font = run.font
        if font_family is None and font.name:
            font_family = font.name
        if font_size is None and font.size is not None:
            font_size = round(float(font.size) * _EMU_TO_PT, 1)
        if font_weight is None and font.bold:
            font_weight = "bold"
        if font_style is None and font.italic:
            font_style = "italic"

    return LayoutHint(
        reading_order=reading_order,
        width=content_width,
        font_family=font_family,
        font_size=font_size,
        font_weight=font_weight,
        font_style=font_style,
    )

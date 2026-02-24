"""Tests for DOCX, PDF, and Google Docs format handlers."""

from __future__ import annotations

import json
from typing import TYPE_CHECKING

from docx import Document

from uaf.app.formats.docx_format import DocxComparator, DocxHandler
from uaf.app.formats.gdoc_format import GdocHandler
from uaf.app.formats.pdf_format import PdfHandler
from uaf.core.nodes import Artifact, Cell, Heading, Paragraph, Sheet
from uaf.db.graph_db import GraphDB

if TYPE_CHECKING:
    from pathlib import Path


# ---------------------------------------------------------------------------
# DOCX tests
# ---------------------------------------------------------------------------


class TestDocxHandler:
    """Tests for DocxHandler import/export."""

    def test_import_paragraphs(self, tmp_path: Path) -> None:
        """Paragraphs in a DOCX are imported as Paragraph nodes."""
        docx_path = tmp_path / "test.docx"
        doc = Document()
        doc.add_paragraph("First paragraph")
        doc.add_paragraph("Second paragraph")
        doc.save(str(docx_path))

        db = GraphDB()
        handler = DocxHandler()
        root_id = handler.import_file(docx_path, db)

        art = db.get_node(root_id)
        assert isinstance(art, Artifact)
        assert art.title == "test"

        children = db.get_children(root_id)
        paragraphs = [c for c in children if isinstance(c, Paragraph)]
        assert len(paragraphs) == 2
        assert paragraphs[0].text == "First paragraph"
        assert paragraphs[1].text == "Second paragraph"

    def test_import_headings(self, tmp_path: Path) -> None:
        """Headings in a DOCX are imported as Heading nodes."""
        docx_path = tmp_path / "headings.docx"
        doc = Document()
        doc.add_heading("Title Heading", level=1)
        doc.add_paragraph("Body text")
        doc.add_heading("Sub Heading", level=2)
        doc.save(str(docx_path))

        db = GraphDB()
        handler = DocxHandler()
        root_id = handler.import_file(docx_path, db)

        children = db.get_children(root_id)
        headings = [c for c in children if isinstance(c, Heading)]
        assert len(headings) == 2
        assert headings[0].text == "Title Heading"
        assert headings[0].level == 1
        assert headings[1].text == "Sub Heading"
        assert headings[1].level == 2

    def test_import_table(self, tmp_path: Path) -> None:
        """Tables in a DOCX are imported as Sheet + Cell nodes."""
        docx_path = tmp_path / "table.docx"
        doc = Document()
        table = doc.add_table(rows=2, cols=3)
        table.rows[0].cells[0].text = "A"
        table.rows[0].cells[1].text = "B"
        table.rows[0].cells[2].text = "C"
        table.rows[1].cells[0].text = "1"
        table.rows[1].cells[1].text = "2"
        table.rows[1].cells[2].text = "3"
        doc.save(str(docx_path))

        db = GraphDB()
        handler = DocxHandler()
        root_id = handler.import_file(docx_path, db)

        children = db.get_children(root_id)
        sheets = [c for c in children if isinstance(c, Sheet)]
        assert len(sheets) == 1
        assert sheets[0].rows == 2
        assert sheets[0].cols == 3

        cells = db.get_children(sheets[0].meta.id)
        cell_nodes = [c for c in cells if isinstance(c, Cell)]
        assert len(cell_nodes) == 6

    def test_export_roundtrip(self, tmp_path: Path) -> None:
        """Export produces a valid DOCX with the same text content."""
        docx_path = tmp_path / "original.docx"
        doc = Document()
        doc.add_heading("My Title", level=1)
        doc.add_paragraph("Hello world")
        doc.save(str(docx_path))

        db = GraphDB()
        handler = DocxHandler()
        root_id = handler.import_file(docx_path, db)

        output = tmp_path / "exported.docx"
        handler.export_file(db, root_id, output)

        result = DocxComparator().compare(docx_path, output)
        assert result.is_equivalent
        assert result.similarity_score >= 0.95

    def test_empty_docx(self, tmp_path: Path) -> None:
        """An empty DOCX imports without errors."""
        docx_path = tmp_path / "empty.docx"
        doc = Document()
        doc.save(str(docx_path))

        db = GraphDB()
        handler = DocxHandler()
        root_id = handler.import_file(docx_path, db)

        art = db.get_node(root_id)
        assert isinstance(art, Artifact)
        children = db.get_children(root_id)
        assert len(children) == 0


# ---------------------------------------------------------------------------
# PDF tests
# ---------------------------------------------------------------------------


class TestPdfHandler:
    """Tests for PdfHandler import."""

    def test_import_pdf(self, tmp_path: Path) -> None:
        """A PDF with text blocks is imported as Paragraph nodes."""
        import fitz

        pdf_path = tmp_path / "test.pdf"
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((72, 72), "First paragraph of text")
        page.insert_text((72, 100), "Second paragraph of text")
        doc.save(str(pdf_path))
        doc.close()

        db = GraphDB()
        handler = PdfHandler()
        root_id = handler.import_file(pdf_path, db)

        art = db.get_node(root_id)
        assert isinstance(art, Artifact)
        assert art.title == "test"

        children = db.get_children(root_id)
        paragraphs = [c for c in children if isinstance(c, Paragraph)]
        assert len(paragraphs) >= 1
        all_text = " ".join(p.text for p in paragraphs)
        assert "First paragraph" in all_text
        assert "Second paragraph" in all_text

    def test_import_multipage_pdf(self, tmp_path: Path) -> None:
        """A multi-page PDF extracts text from all pages."""
        import fitz

        pdf_path = tmp_path / "multipage.pdf"
        doc = fitz.open()
        for i in range(3):
            page = doc.new_page()
            page.insert_text((72, 72), f"Page {i + 1} content")
        doc.save(str(pdf_path))
        doc.close()

        db = GraphDB()
        handler = PdfHandler()
        root_id = handler.import_file(pdf_path, db)

        children = db.get_children(root_id)
        paragraphs = [c for c in children if isinstance(c, Paragraph)]
        all_text = " ".join(p.text for p in paragraphs)
        assert "Page 1" in all_text
        assert "Page 2" in all_text
        assert "Page 3" in all_text

    def test_export_as_text(self, tmp_path: Path) -> None:
        """PDF handler exports as plain text."""
        import fitz

        pdf_path = tmp_path / "test.pdf"
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((72, 72), "Hello from PDF")
        doc.save(str(pdf_path))
        doc.close()

        db = GraphDB()
        handler = PdfHandler()
        root_id = handler.import_file(pdf_path, db)

        output = tmp_path / "output.txt"
        handler.export_file(db, root_id, output)

        text = output.read_text(encoding="utf-8")
        assert "Hello from PDF" in text


    def test_extract_text_no_extra_spaces(self, tmp_path: Path) -> None:
        """Small-caps style spans don't produce extra spaces in text."""
        import fitz

        pdf_path = tmp_path / "caps.pdf"
        doc = fitz.open()
        page = doc.new_page()
        # Simulate small-caps by inserting text at two sizes on the same line.
        page.insert_text((72, 72), "D", fontsize=17)
        page.insert_text((83, 72), "YNAMIC ", fontsize=14)
        page.insert_text((140, 72), "N", fontsize=17)
        page.insert_text((152, 72), "ESTED", fontsize=14)
        doc.save(str(pdf_path))
        doc.close()

        db = GraphDB()
        handler = PdfHandler()
        root_id = handler.import_file(pdf_path, db)

        children = db.get_children(root_id)
        paragraphs = [c for c in children if isinstance(c, Paragraph)]
        all_text = " ".join(p.text for p in paragraphs)
        # Must NOT contain double spaces.
        assert "  " not in all_text

    def test_extract_text_preserves_line_breaks(self, tmp_path: Path) -> None:
        """Multi-line text blocks preserve newlines between lines."""
        import fitz

        pdf_path = tmp_path / "multiline.pdf"
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((72, 72), "First line of text")
        page.insert_text((72, 86), "Second line of text")
        doc.save(str(pdf_path))
        doc.close()

        db = GraphDB()
        handler = PdfHandler()
        root_id = handler.import_file(pdf_path, db)

        children = db.get_children(root_id)
        paragraphs = [c for c in children if isinstance(c, Paragraph)]
        # At least one paragraph should contain a newline if both lines
        # ended up in the same block, OR the two lines should be separate
        # paragraphs with clean text.
        all_text = " ".join(p.text for p in paragraphs)
        assert "First line" in all_text
        assert "Second line" in all_text

    def test_dehyphenation(self, tmp_path: Path) -> None:
        """End-of-line hyphens splitting words are removed during import."""
        from uaf.app.formats.pdf_format import _extract_block_text

        # Simulate a block with two lines: "capa-" and "bilities"
        block: dict[str, object] = {
            "lines": [
                {
                    "spans": [{"text": "remarkable capa-"}],
                    "dir": (1, 0),
                },
                {
                    "spans": [{"text": "bilities in various"}],
                    "dir": (1, 0),
                },
            ],
        }
        text = _extract_block_text(block)
        assert "capabilities" in text
        assert "capa-" not in text

    def test_dehyphenation_preserves_real_hyphens(self, tmp_path: Path) -> None:
        """Hyphens that are not line-end word splits are preserved."""
        from uaf.app.formats.pdf_format import _extract_block_text

        block: dict[str, object] = {
            "lines": [
                {
                    "spans": [{"text": "well-known method"}],
                    "dir": (1, 0),
                },
            ],
        }
        text = _extract_block_text(block)
        assert "well-known" in text

    def test_bold_detection_first_line(self, tmp_path: Path) -> None:
        """Bold from the first line is preserved even if later lines are normal."""
        from uaf.app.formats.pdf_format import _extract_dominant_font

        # First line bold (flags bit 4 = 16), second line normal
        block: dict[str, object] = {
            "lines": [
                {
                    "spans": [
                        {"text": "Author Name", "font": "Helvetica-Bold",
                         "size": 10.0, "flags": 16, "color": 0},
                    ],
                },
                {
                    "spans": [
                        {"text": "University of Something and more text here",
                         "font": "Helvetica", "size": 10.0, "flags": 0,
                         "color": 0},
                    ],
                },
            ],
        }
        font = _extract_dominant_font(block)
        assert font.get("weight") == "bold"

    def test_rotation_extraction(self, tmp_path: Path) -> None:
        """Rotated text blocks have rotation stored in LayoutHint."""
        from uaf.app.formats.pdf_format import _extract_rotation

        # dir = (0, -1) → text reads bottom-to-top → -90°
        block: dict[str, object] = {
            "lines": [
                {"dir": (0.0, -1.0), "spans": [{"text": "Sidebar"}]},
            ],
        }
        angle = _extract_rotation(block)
        assert angle is not None
        assert abs(angle - (-90.0)) < 0.5

    def test_horizontal_text_no_rotation(self, tmp_path: Path) -> None:
        """Horizontal text (dir=(1,0)) yields None rotation."""
        from uaf.app.formats.pdf_format import _extract_rotation

        block: dict[str, object] = {
            "lines": [
                {"dir": (1.0, 0.0), "spans": [{"text": "Normal"}]},
            ],
        }
        angle = _extract_rotation(block)
        assert angle is None


# ---------------------------------------------------------------------------
# Google Docs JSON tests
# ---------------------------------------------------------------------------


class TestGdocHandler:
    """Tests for GdocHandler import/export."""

    def _make_gdoc_json(
        self,
        title: str,
        elements: list[dict[str, object]],
    ) -> dict[str, object]:
        """Build a minimal Google Docs JSON structure."""
        return {
            "title": title,
            "body": {"content": elements},
        }

    def _make_paragraph_element(
        self, text: str, style: str = "NORMAL_TEXT",
    ) -> dict[str, object]:
        return {
            "paragraph": {
                "paragraphStyle": {"namedStyleType": style},
                "elements": [{"textRun": {"content": text + "\n"}}],
            },
        }

    def test_import_paragraphs(self, tmp_path: Path) -> None:
        """Normal text paragraphs are imported as Paragraph nodes."""
        data = self._make_gdoc_json("Test Doc", [
            self._make_paragraph_element("First paragraph"),
            self._make_paragraph_element("Second paragraph"),
        ])

        gdoc_path = tmp_path / "test.json"
        gdoc_path.write_text(json.dumps(data), encoding="utf-8")

        db = GraphDB()
        handler = GdocHandler()
        root_id = handler.import_file(gdoc_path, db)

        art = db.get_node(root_id)
        assert isinstance(art, Artifact)
        assert art.title == "Test Doc"

        children = db.get_children(root_id)
        paragraphs = [c for c in children if isinstance(c, Paragraph)]
        assert len(paragraphs) == 2
        assert paragraphs[0].text == "First paragraph"

    def test_import_headings(self, tmp_path: Path) -> None:
        """HEADING_N styled paragraphs are imported as Heading nodes."""
        data = self._make_gdoc_json("Heading Test", [
            self._make_paragraph_element("Main Title", "HEADING_1"),
            self._make_paragraph_element("Body text"),
            self._make_paragraph_element("Subtitle", "HEADING_2"),
        ])

        gdoc_path = tmp_path / "headings.json"
        gdoc_path.write_text(json.dumps(data), encoding="utf-8")

        db = GraphDB()
        handler = GdocHandler()
        root_id = handler.import_file(gdoc_path, db)

        children = db.get_children(root_id)
        headings = [c for c in children if isinstance(c, Heading)]
        assert len(headings) == 2
        assert headings[0].text == "Main Title"
        assert headings[0].level == 1
        assert headings[1].text == "Subtitle"
        assert headings[1].level == 2

    def test_export_roundtrip(self, tmp_path: Path) -> None:
        """Export and re-import produces the same content."""
        data = self._make_gdoc_json("Roundtrip", [
            self._make_paragraph_element("Title", "HEADING_1"),
            self._make_paragraph_element("Some body text"),
        ])

        gdoc_path = tmp_path / "original.json"
        gdoc_path.write_text(json.dumps(data), encoding="utf-8")

        db = GraphDB()
        handler = GdocHandler()
        root_id = handler.import_file(gdoc_path, db)

        output = tmp_path / "exported.json"
        handler.export_file(db, root_id, output)

        # Re-import the exported file
        db2 = GraphDB()
        root_id2 = handler.import_file(output, db2)
        children2 = db2.get_children(root_id2)

        headings = [c for c in children2 if isinstance(c, Heading)]
        paragraphs = [c for c in children2 if isinstance(c, Paragraph)]
        assert len(headings) == 1
        assert headings[0].text == "Title"
        assert len(paragraphs) == 1
        assert paragraphs[0].text == "Some body text"

    def test_empty_body(self, tmp_path: Path) -> None:
        """A Google Docs JSON with empty body imports without errors."""
        data = {"title": "Empty", "body": {"content": []}}

        gdoc_path = tmp_path / "empty.json"
        gdoc_path.write_text(json.dumps(data), encoding="utf-8")

        db = GraphDB()
        handler = GdocHandler()
        root_id = handler.import_file(gdoc_path, db)

        art = db.get_node(root_id)
        assert isinstance(art, Artifact)
        assert art.title == "Empty"
        children = db.get_children(root_id)
        assert len(children) == 0

    def test_skips_empty_paragraphs(self, tmp_path: Path) -> None:
        """Paragraphs with only whitespace are skipped."""
        data = self._make_gdoc_json("Sparse", [
            self._make_paragraph_element("Real content"),
            {"paragraph": {
                "paragraphStyle": {"namedStyleType": "NORMAL_TEXT"},
                "elements": [{"textRun": {"content": "   \n"}}],
            }},
            self._make_paragraph_element("More content"),
        ])

        gdoc_path = tmp_path / "sparse.json"
        gdoc_path.write_text(json.dumps(data), encoding="utf-8")

        db = GraphDB()
        handler = GdocHandler()
        root_id = handler.import_file(gdoc_path, db)

        children = db.get_children(root_id)
        paragraphs = [c for c in children if isinstance(c, Paragraph)]
        assert len(paragraphs) == 2

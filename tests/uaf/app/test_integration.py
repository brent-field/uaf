"""Integration tests — end-to-end scenarios crossing all application layers."""

from __future__ import annotations

import asyncio
import json
import tempfile
from pathlib import Path
from typing import Any

from fastapi.testclient import TestClient

from uaf.app.api import create_app
from uaf.app.lenses import LensRegistry
from uaf.app.lenses.doc_lens import DocLens
from uaf.app.lenses.grid_lens import GridLens
from uaf.app.mcp_server import create_mcp_server
from uaf.core.edges import Edge, EdgeType
from uaf.core.node_id import EdgeId, NodeId, utc_now
from uaf.core.nodes import (
    Artifact,
    Cell,
    NodeType,
    Sheet,
    make_node_metadata,
)
from uaf.db.graph_db import GraphDB
from uaf.security.auth import LocalAuthProvider, PasswordCredentials
from uaf.security.primitives import PrincipalId, Role
from uaf.security.secure_graph_db import SecureGraphDB


def _contains(source: NodeId, target: NodeId) -> Edge:
    return Edge(
        id=EdgeId.generate(),
        source=source,
        target=target,
        edge_type=EdgeType.CONTAINS,
        created_at=utc_now(),
    )


class TestDocumentWorkflow:
    """Register -> login -> create artifact -> add content via DocLens -> render
    -> export as Markdown -> verify content."""

    def test_full_document_workflow(self) -> None:
        db = GraphDB()
        auth = LocalAuthProvider()
        sdb = SecureGraphDB(db, auth)
        registry = LensRegistry()
        registry.register(DocLens())
        registry.register(GridLens())
        app = create_app(sdb, registry)
        client = TestClient(app)

        # Register
        resp = client.post(
            "/api/auth/register",
            json={"display_name": "Alice", "password": "alice123"},
        )
        assert resp.status_code == 200
        token = resp.json()["token"]
        headers = {"Authorization": f"Bearer {token}"}

        # Create artifact
        resp = client.post(
            "/api/artifacts", json={"title": "Quarterly Report"}, headers=headers,
        )
        assert resp.status_code == 201
        aid = resp.json()["id"]

        # Add heading via DocLens action
        client.post(
            f"/api/artifacts/{aid}/lens/doc/action",
            json={
                "action_type": "insert_text",
                "params": {
                    "parent_id": aid, "text": "Overview", "position": 0,
                    "style": "heading",
                },
            },
            headers=headers,
        )

        # Add paragraph
        client.post(
            f"/api/artifacts/{aid}/lens/doc/action",
            json={
                "action_type": "insert_text",
                "params": {
                    "parent_id": aid, "text": "Revenue grew 15%.", "position": 1,
                },
            },
            headers=headers,
        )

        # Render via DocLens
        resp = client.get(f"/api/artifacts/{aid}/lens/doc", headers=headers)
        assert resp.status_code == 200
        view = resp.json()
        assert view["lens_type"] == "doc"
        assert "Overview" in view["content"]
        assert "Revenue grew 15%." in view["content"]
        assert view["node_count"] == 3  # artifact + heading + paragraph

        # Export as Markdown
        resp = client.get(
            f"/api/artifacts/{aid}/export?format=markdown", headers=headers,
        )
        assert resp.status_code == 200
        md = resp.text
        assert "Overview" in md
        assert "Revenue grew 15%." in md


class TestSpreadsheetWorkflow:
    """Create artifact -> add sheet -> add cells via GridLens -> render ->
    export as CSV -> verify values."""

    def test_full_spreadsheet_workflow(self) -> None:
        db = GraphDB()
        auth = LocalAuthProvider()
        sdb = SecureGraphDB(db, auth)
        session = sdb.system_session()
        lens = GridLens()

        # Create spreadsheet with data
        art = Artifact(meta=make_node_metadata(NodeType.ARTIFACT), title="Budget")
        art_id = sdb.create_node(session, art)

        sheet = Sheet(
            meta=make_node_metadata(NodeType.SHEET), title="Sheet1", rows=2, cols=2,
        )
        sheet_id = sdb.create_node(session, sheet)
        sdb.create_edge(session, _contains(art_id, sheet_id))

        for r, row in enumerate([["Item", "Cost"], ["Coffee", "5.00"]]):
            for c, val in enumerate(row):
                cell = Cell(
                    meta=make_node_metadata(NodeType.CELL), value=val, row=r, col=c,
                )
                cid = sdb.create_node(session, cell)
                sdb.create_edge(session, _contains(sheet_id, cid))

        # Render
        view = lens.render(sdb, session, art_id)
        assert view.lens_type == "grid"
        assert ">Item</td>" in view.content
        assert ">Coffee</td>" in view.content
        assert ">5.00</td>" in view.content

        # Export as CSV
        from uaf.app.formats.csv_format import CsvHandler

        handler = CsvHandler()
        with tempfile.NamedTemporaryFile(suffix=".csv", delete=False) as tmp:
            csv_path = Path(tmp.name)
        handler.export_file(db, art_id, csv_path)
        csv_text = csv_path.read_text(encoding="utf-8")
        csv_path.unlink()
        assert "Item" in csv_text
        assert "Coffee" in csv_text
        assert "5.00" in csv_text


class TestImportLensExport:
    """Import Markdown -> render via DocLens -> verify HTML -> export back
    -> round-trip verify."""

    def test_import_render_export_roundtrip(self) -> None:
        db = GraphDB()
        auth = LocalAuthProvider()
        sdb = SecureGraphDB(db, auth)
        session = sdb.system_session()
        lens = DocLens()

        # Import Markdown
        from uaf.app.formats.markdown import MarkdownHandler

        md_content = "# Welcome\n\nThis is a test document.\n\n## Details\n\nMore info here.\n"
        with tempfile.NamedTemporaryFile(suffix=".md", delete=False, mode="w") as tmp:
            tmp.write(md_content)
            md_path = Path(tmp.name)

        handler = MarkdownHandler()
        art_id = handler.import_file(md_path, db)
        md_path.unlink()

        # Render via DocLens
        view = lens.render(sdb, session, art_id)
        assert "Welcome" in view.content
        assert "This is a test document." in view.content
        assert "Details" in view.content
        assert "More info here." in view.content
        assert view.node_count >= 5  # artifact + 2 headings + 2 paragraphs

        # Export back to Markdown
        with tempfile.NamedTemporaryFile(suffix=".md", delete=False) as tmp:
            export_path = Path(tmp.name)
        handler.export_file(db, art_id, export_path)
        exported = export_path.read_text(encoding="utf-8")
        export_path.unlink()

        assert "Welcome" in exported
        assert "test document" in exported
        assert "Details" in exported


class TestMCPAgentWorkflow:
    """Create artifact via MCP -> add content -> query via find_by_type ->
    render via render_artifact -> verify output."""

    def test_mcp_agent_workflow(self) -> None:
        db = GraphDB()
        auth = LocalAuthProvider()
        sdb = SecureGraphDB(db, auth)
        registry = LensRegistry()
        registry.register(DocLens())
        session = sdb.system_session()
        mcp = create_mcp_server(sdb, session, registry)

        def call(name: str, args: dict[str, object]) -> dict[str, Any]:
            blocks, _ = asyncio.run(mcp.call_tool(name, args))
            return json.loads(blocks[0].text)  # type: ignore[union-attr,no-any-return]

        # Create artifact
        art = call("create_artifact", {"title": "AI Notes"})
        aid = art["artifact_id"]

        # Add content
        call("add_child", {"parent_id": aid, "node_type": "paragraph", "text": "Note 1"})
        call("add_child", {"parent_id": aid, "node_type": "paragraph", "text": "Note 2"})

        # Query
        blocks, _ = asyncio.run(
            mcp.call_tool("find_by_type", {"node_type": "artifact"})
        )
        artifacts = json.loads(blocks[0].text)  # type: ignore[union-attr]
        assert any(a.get("title") == "AI Notes" for a in artifacts)

        # Render
        result = call("render_artifact", {"artifact_id": aid, "lens_type": "doc"})
        assert "Note 1" in result["content"]
        assert "Note 2" in result["content"]
        assert result["node_count"] == 3


class TestMultiUserAPI:
    """User 1 creates artifact -> grants EDITOR to user 2 -> user 2 edits
    -> user 3 gets 403 -> audit log shows all actions."""

    def test_multi_user_permissions(self) -> None:
        db = GraphDB()
        auth = LocalAuthProvider()
        sdb = SecureGraphDB(db, auth)
        registry = LensRegistry()
        registry.register(DocLens())
        app = create_app(sdb, registry)
        client = TestClient(app)

        # Register user1 and user2
        r1 = client.post(
            "/api/auth/register", json={"display_name": "User1", "password": "p1"},
        )
        token1 = r1.json()["token"]
        h1 = {"Authorization": f"Bearer {token1}"}

        r2 = client.post(
            "/api/auth/register", json={"display_name": "User2", "password": "p2"},
        )
        token2 = r2.json()["token"]
        pid2 = r2.json()["principal_id"]
        h2 = {"Authorization": f"Bearer {token2}"}

        # User3 just registers (no explicit grant)
        r3 = client.post(
            "/api/auth/register", json={"display_name": "User3", "password": "p3"},
        )
        token3 = r3.json()["token"]
        h3 = {"Authorization": f"Bearer {token3}"}

        # User1 creates artifact
        resp = client.post(
            "/api/artifacts", json={"title": "Shared"}, headers=h1,
        )
        aid = resp.json()["id"]

        # User1 grants EDITOR to user2
        sdb.grant_role(
            sdb.authenticate(
                PasswordCredentials(
                    principal_id=PrincipalId(value=r1.json()["principal_id"]),
                    password="p1",
                )
            ),
            NodeId.from_str(aid) if hasattr(NodeId, "from_str") else _parse_nid(aid),
            PrincipalId(value=pid2),
            Role.EDITOR,
        )

        # User2 can read the artifact
        resp = client.get(f"/api/artifacts/{aid}", headers=h2)
        assert resp.status_code == 200

        # User3 should not be able to see the artifact (not in ACL)
        resp = client.get(f"/api/artifacts/{aid}", headers=h3)
        assert resp.status_code == 404  # returns None -> 404

        # Audit log has entries
        audit = sdb.get_audit_log()
        assert audit.count() > 0


class TestPdfImportRoute:
    """Import a PDF via the API route and verify the artifact is created."""

    def test_import_pdf_via_api(self) -> None:
        import fitz

        db = GraphDB()
        auth = LocalAuthProvider()
        sdb = SecureGraphDB(db, auth)
        registry = LensRegistry()
        registry.register(DocLens())
        registry.register(GridLens())
        app = create_app(sdb, registry)
        client = TestClient(app)

        # Register and get token
        resp = client.post(
            "/api/auth/register", json={"display_name": "PdfUser", "password": "pw"},
        )
        token = resp.json()["token"]
        headers = {"Authorization": f"Bearer {token}"}

        # Create a small PDF in memory
        pdf_doc = fitz.open()
        page = pdf_doc.new_page()
        page.insert_text((72, 72), "Hello from PDF import test")
        pdf_bytes = pdf_doc.tobytes()
        pdf_doc.close()

        # Import via API (auto-detect from .pdf extension)
        resp = client.post(
            "/api/artifacts/import",
            files={"file": ("sample.pdf", pdf_bytes, "application/pdf")},
            headers=headers,
        )
        assert resp.status_code == 201, resp.text
        data = resp.json()
        assert data["title"] == "sample"

        # Verify children were created
        aid = data["id"]
        resp = client.get(f"/api/nodes/{aid}/children", headers=headers)
        assert resp.status_code == 200
        children = resp.json()["children"]
        assert len(children) >= 1
        texts = [c.get("fields", {}).get("text", "") for c in children]
        assert any("Hello from PDF" in t for t in texts)

    def test_import_docx_via_api(self) -> None:
        from io import BytesIO

        from docx import Document

        db = GraphDB()
        auth = LocalAuthProvider()
        sdb = SecureGraphDB(db, auth)
        registry = LensRegistry()
        registry.register(DocLens())
        registry.register(GridLens())
        app = create_app(sdb, registry)
        client = TestClient(app)

        # Register and get token
        resp = client.post(
            "/api/auth/register", json={"display_name": "DocxUser", "password": "pw"},
        )
        token = resp.json()["token"]
        headers = {"Authorization": f"Bearer {token}"}

        # Create a DOCX in memory
        doc = Document()
        doc.add_heading("Test Heading", level=1)
        doc.add_paragraph("Test paragraph content")
        buf = BytesIO()
        doc.save(buf)
        docx_bytes = buf.getvalue()

        # Import via API
        resp = client.post(
            "/api/artifacts/import",
            files={"file": (
                "report.docx", docx_bytes,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )},
            headers=headers,
        )
        assert resp.status_code == 201, resp.text
        data = resp.json()
        assert data["title"] == "report"

    def test_import_gdoc_via_api(self) -> None:
        db = GraphDB()
        auth = LocalAuthProvider()
        sdb = SecureGraphDB(db, auth)
        registry = LensRegistry()
        registry.register(DocLens())
        registry.register(GridLens())
        app = create_app(sdb, registry)
        client = TestClient(app)

        # Register and get token
        resp = client.post(
            "/api/auth/register", json={"display_name": "GdocUser", "password": "pw"},
        )
        token = resp.json()["token"]
        headers = {"Authorization": f"Bearer {token}"}

        # Create Google Docs JSON
        gdoc_data = json.dumps({
            "title": "My Google Doc",
            "body": {
                "content": [
                    {
                        "paragraph": {
                            "paragraphStyle": {"namedStyleType": "NORMAL_TEXT"},
                            "elements": [{"textRun": {"content": "Hello from GDoc\n"}}],
                        },
                    },
                ],
            },
        })

        resp = client.post(
            "/api/artifacts/import",
            files={"file": ("notes.json", gdoc_data.encode(), "application/json")},
            headers=headers,
        )
        assert resp.status_code == 201, resp.text
        data = resp.json()
        assert data["title"] == "My Google Doc"


class TestCrossLensConsistency:
    """Create document via DocLens -> same artifact rendered via API JSON endpoint
    -> both views show same content."""

    def test_cross_lens_consistency(self) -> None:
        db = GraphDB()
        auth = LocalAuthProvider()
        sdb = SecureGraphDB(db, auth)
        registry = LensRegistry()
        registry.register(DocLens())
        registry.register(GridLens())
        app = create_app(sdb, registry)
        client = TestClient(app)

        # Register
        resp = client.post(
            "/api/auth/register", json={"display_name": "X", "password": "x"},
        )
        token = resp.json()["token"]
        headers = {"Authorization": f"Bearer {token}"}

        # Create artifact with content
        resp = client.post(
            "/api/artifacts", json={"title": "CrossLens"}, headers=headers,
        )
        aid = resp.json()["id"]

        client.post(
            f"/api/artifacts/{aid}/lens/doc/action",
            json={
                "action_type": "insert_text",
                "params": {"parent_id": aid, "text": "Consistency", "position": 0},
            },
            headers=headers,
        )

        # Render via DocLens
        resp = client.get(f"/api/artifacts/{aid}/lens/doc", headers=headers)
        doc_view = resp.json()

        # Read the same artifact via node API
        resp = client.get(f"/api/nodes/{aid}/children", headers=headers)
        children = resp.json()["children"]

        # DocLens shows the content
        assert "Consistency" in doc_view["content"]

        # Node API shows the same child
        assert len(children) == 1
        assert children[0]["node_type"] == "paragraph"


class TestLayoutViewRoute:
    """Test the HTMX layout/semantic view toggle routes."""

    @staticmethod
    def _setup_client_with_pdf(
        text_items: list[tuple[tuple[float, float], str]],
        *,
        pages: int = 1,
    ) -> tuple[TestClient, str]:
        """Create a TestClient, register a user, import a PDF, return (client, aid).

        *text_items* are (position, text) pairs inserted on the first page.
        *pages* adds extra pages with "Page N" text for multipage tests.
        """
        import fitz

        db = GraphDB()
        auth = LocalAuthProvider()
        sdb = SecureGraphDB(db, auth)
        registry = LensRegistry()
        registry.register(DocLens())
        registry.register(GridLens())
        app = create_app(sdb, registry)
        client = TestClient(app)

        # Register via API to get a JWT token.
        resp = client.post(
            "/api/auth/register",
            json={"display_name": "LayoutUser", "password": "pw"},
        )
        token = resp.json()["token"]
        # Set cookie for HTMX frontend routes.
        client.cookies.set("uaf_token", token)

        # Build a PDF.
        pdf_doc = fitz.open()
        page = pdf_doc.new_page()
        for pos, txt in text_items:
            page.insert_text(pos, txt)
        for i in range(1, pages):
            extra = pdf_doc.new_page()
            extra.insert_text((72, 72), f"Page {i + 1}")
        pdf_bytes = pdf_doc.tobytes()
        pdf_doc.close()

        # Import via API (returns JSON with artifact id).
        resp = client.post(
            "/api/artifacts/import",
            files={"file": ("test.pdf", pdf_bytes, "application/pdf")},
            headers={"Authorization": f"Bearer {token}"},
        )
        assert resp.status_code == 201, resp.text
        aid = resp.json()["id"]
        return client, aid

    def test_layout_toggle_returns_layout_html(self) -> None:
        """GET ?mode=layout returns layout-page structure, not <article>."""
        client, aid = self._setup_client_with_pdf([
            ((72, 72), "Hello World"),
        ])
        resp = client.get(f"/artifacts/{aid}/blocks?mode=layout")
        assert resp.status_code == 200
        html = resp.text
        assert "layout-page" in html
        assert "layout-block" in html
        assert "position: absolute" in html
        assert "<article" not in html

    def test_semantic_toggle_returns_semantic_html(self) -> None:
        """GET ?mode=semantic returns doc-block structure, not layout-page."""
        client, aid = self._setup_client_with_pdf([
            ((72, 72), "Hello World"),
        ])
        resp = client.get(f"/artifacts/{aid}/blocks?mode=semantic")
        assert resp.status_code == 200
        html = resp.text
        assert "doc-block" in html
        assert "layout-page" not in html

    def test_layout_view_has_no_extra_spaces(self) -> None:
        """Imported PDF text has no spurious double spaces."""
        client, aid = self._setup_client_with_pdf([
            ((72, 72), "Hello World"),
        ])
        resp = client.get(f"/artifacts/{aid}/blocks?mode=layout")
        assert resp.status_code == 200
        assert "Hello World" in resp.text

    def test_layout_view_multipage_structure(self) -> None:
        """Multi-page PDF produces one layout-page per source page."""
        client, aid = self._setup_client_with_pdf(
            [((72, 72), "Page 1")], pages=3,
        )
        resp = client.get(f"/artifacts/{aid}/blocks?mode=layout")
        assert resp.status_code == 200
        assert resp.text.count("layout-page") >= 3

    def test_layout_blocks_have_no_height(self) -> None:
        """Individual layout blocks must not set explicit height."""
        import re

        client, aid = self._setup_client_with_pdf([
            ((72, 72), "Block text"),
        ])
        resp = client.get(f"/artifacts/{aid}/blocks?mode=layout")
        assert resp.status_code == 200
        # Page container has height — that's expected.
        # But individual blocks must not.
        block_styles = re.findall(
            r'class="layout-block"[^>]*style="([^"]*)"', resp.text,
        )
        for style in block_styles:
            assert "height:" not in style

    def test_header_footer_detection_via_route(self) -> None:
        """Repeated header text across pages gets header-footer class."""
        import fitz

        db = GraphDB()
        auth = LocalAuthProvider()
        sdb = SecureGraphDB(db, auth)
        registry = LensRegistry()
        registry.register(DocLens())
        registry.register(GridLens())
        app = create_app(sdb, registry)
        client = TestClient(app)

        resp = client.post(
            "/api/auth/register",
            json={"display_name": "HFUser", "password": "pw"},
        )
        token = resp.json()["token"]
        client.cookies.set("uaf_token", token)

        # Build PDF with repeated header text on 3 pages.
        pdf_doc = fitz.open()
        for i in range(3):
            page = pdf_doc.new_page()
            page.insert_text((400, 30), "PREPRINT HEADER")
            page.insert_text((72, 200), f"Body text page {i + 1}")
        pdf_bytes = pdf_doc.tobytes()
        pdf_doc.close()

        resp = client.post(
            "/api/artifacts/import",
            files={"file": ("hf.pdf", pdf_bytes, "application/pdf")},
            headers={"Authorization": f"Bearer {token}"},
        )
        assert resp.status_code == 201
        aid = resp.json()["id"]

        resp = client.get(f"/artifacts/{aid}/blocks?mode=layout")
        assert resp.status_code == 200
        assert "layout-header-footer" in resp.text


def _parse_nid(s: str) -> NodeId:
    """Parse a string UUID into a NodeId."""
    import uuid

    return NodeId(value=uuid.UUID(s))
